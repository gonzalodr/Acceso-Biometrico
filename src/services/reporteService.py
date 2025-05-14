import re
import csv
import os

from models.reporte     import Reporte
from data.reporteData   import ReporteData
from PySide6.QtCore     import *
from fpdf               import FPDF
from docx               import Document
from collections        import defaultdict
import datetime


class ReporteServices:
    def __init__(self):
        self.reporteData = ReporteData()
        
    def insertarReporte(self, reporte: Reporte):
        reporte.fecha_generacion = QDate.currentDate().toString("yyyy-MM-dd")
        return self.reporteData.create_reporte(reporte)
    
    def modificarReporte(self, reporte: Reporte):
        return self.reporteData.update_reporte(reporte)
        
    def eliminarReporte(self, id):
        return self.reporteData.delete_reporte(id)

    def obtenerListaReporte(self, pagina=1, tam_pagina=10, ordenar_por="id", tipo_orden="ASC", busqueda=None):
        return self.reporteData.list_reportes(pagina, tam_pagina, ordenar_por, tipo_orden, busqueda)

    def obtenerReportePorId(self, id):
        return self.reporteData.get_reporte_by_id(id)
    
    def obtenerTodoReporte(self):
        return self.reporteData.obtener_todo_reportes()
    
    def __obtener_datos_en_lotes(self, limit=500, **filtros):
            offset = 0
            while True:
                resultado = self.reporteData.obtener_datos_para_reporte(limit=limit, offset=offset, **filtros)
                if not resultado['success']:
                    yield {'success': False, 'message': resultado['message']}
                    break
                
                lote = resultado['reporte']
                if not lote:
                    break
                yield lote
                offset += limit

    def crear_reporte(self,ruta:str,nombre_reporte:str, extencion:str,**filtros):

        rutaFinal = os.path.join(ruta,nombre_reporte+'.'+extencion)

        resultado = self.reporteData.obtener_datos_para_reporte(limit=1, offset=0,**filtros.copy())
        if not resultado['success']:
            return resultado
        
        if resultado['reporte']:
            reporte = resultado['reporte']
            if not any(r in reporte[0] for r in ['justificacion', 'asistencia', 'permiso']):
                return {'success': False, 'message': 'No se encontraron datos.'}
        
        if extencion == 'docx':
           return self.__generar_reporte_docx_por_lotes(rutaFichero=rutaFinal,**filtros)
        elif extencion == 'pdf':
           return self.__generar_reporte_pdf_por_lotes(rutaFichero=rutaFinal,**filtros)
        elif extencion == 'csv':
           return self.__generar_reporte_csv_por_lotes(rutaFichero=rutaFinal,**filtros)
        else:
            return {'success':False,'message':'La extension no es valida'}
    
    #crea el reporte en formato pdf
    def __generar_reporte_pdf_por_lotes(self, rutaFichero="reporte.pdf", limit=1000, **filtros):
        tipoReporte = filtros.get('tipoReporte')
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=10)
        pdf.set_font("Arial", size=9)
        

        # Diccionario para agrupar los datos por cédula de empleado
        empleados_data = {}

        for lote in self.__obtener_datos_en_lotes(limit=limit, **filtros):
            if isinstance(lote, dict) and 'success' in lote:
                return lote

            for empleado in lote:
                cedula = empleado['cedula']
                if cedula not in empleados_data:
                    empleados_data[cedula] = {
                        'id_empleado': empleado['id_empleado'],
                        'nombre': empleado['nombre'],
                        'apellidos': empleado['apellidos'],
                        'cedula':cedula,
                        'departamento': empleado.get('departamento'),
                        'rol': empleado.get('rol'),
                        'justificaciones': [],
                        'asistencias': [],
                        'permisos': []
                    }

                if 'justificacion' in empleado:
                    empleados_data[cedula]['justificaciones'].append(empleado['justificacion'])

                if 'asistencia' in empleado:
                    empleados_data[cedula]['asistencias'].append(empleado['asistencia'])

                if 'permisos' in empleado:
                    empleados_data[cedula]['permisos'].append(empleado['permisos'])

        # Ahora que tenemos agrupados los datos por empleado, generamos el PDF
        pdf.add_page()
        for cedula, info in empleados_data.items():
            nombre_completo = f"Nombre: {info['nombre']} {info['apellidos']} - Cedula: {info['apellidos']} "
            pdf.set_font("Arial", style='B', size=10)
            pdf.cell(190, 6, txt=f"{nombre_completo} - {cedula}", ln=True)
            
            reporte:Reporte = Reporte(id_empleado=info['id_empleado'],fecha_generacion=datetime.datetime.now(),tipo_reporte=','.join(tipoReporte),contenido=','.join(tipoReporte))
            resultHistoria = self.insertarReporte(reporte)
            
            if info.get('departamento'):
                pdf.set_font("Arial", size=9)
                dep = info['departamento']
                pdf.cell(190, 5, txt=f"Departamento: {dep.nombre}", ln=True)

            if info.get('rol'):
                rol = info['rol']
                pdf.cell(190, 5, txt=f"Rol: {rol.nombre}", ln=True)

            pdf.set_font("Arial", style='B', size=9)
            if info['justificaciones']:
                pdf.cell(190, 5, txt="Justificaciones:", ln=True)
                pdf.set_font("Arial", size=9)
                for j in info['justificaciones']:
                    pdf.cell(190, 4, txt=f"{j.fecha} - {j.motivo} - {j.descripcion}", ln=True)

            if info['asistencias']:
                pdf.set_font("Arial", style='B', size=9)
                pdf.cell(190, 5, txt="Asistencias:", ln=True)
                pdf.set_font("Arial", size=9)
                for a in info['asistencias']:
                    pdf.cell(190, 4, txt=f"{a.fecha} - {a.estado_asistencia}", ln=True)

            if info['permisos']:
                pdf.set_font("Arial", style='B', size=9)
                pdf.cell(190, 5, txt="Permisos:", ln=True)
                pdf.set_font("Arial", size=9)
                for p in info['permisos']:
                    pdf.cell(190, 4, txt=f"{p.fecha_inicio} a {p.fecha_fin} - {p.tipo}", ln=True)

            pdf.ln(3)  # Espacio pequeño entre empleados

        pdf.output(rutaFichero)
        os.startfile(rutaFichero)
        return {'success': True, 'message': 'Se creó el reporte exitosamente.'}
    
    #generar el docx
    def __generar_reporte_docx_por_lotes(self, rutaFichero="reporte.docx", limit=1000, **filtros):
        tipoReporte = filtros.get('tipoReporte')
        doc = Document()
        doc.add_heading("Reporte de Empleados", 0)

        empleados_dict = {}

        for lote in self.__obtener_datos_en_lotes(limit=limit, **filtros):
            if isinstance(lote, dict) and 'success' in lote:
                return lote

            for empleado in lote:
                cedula = empleado['cedula']

                if cedula not in empleados_dict:
                    empleados_dict[cedula] = {
                        'id_empleado': empleado['id_empleado'],
                        'nombre': empleado['nombre'],
                        'apellidos': empleado['apellidos'],
                        'cedula': cedula,
                        'departamento': empleado.get('departamento'),
                        'rol': empleado.get('rol'),
                        'justificaciones': [],
                        'asistencias': [],
                        'permisos': []
                    }

                if 'justificacion' in empleado:
                    empleados_dict[cedula]['justificaciones'].append(empleado['justificacion'])
                if 'asistencia' in empleado:
                    empleados_dict[cedula]['asistencias'].append(empleado['asistencia'])
                if 'permisos' in empleado:
                    empleados_dict[cedula]['permisos'].append(empleado['permisos'])

        for emp in empleados_dict.values():
            doc.add_heading(f"Nombre: {emp['nombre']} {emp['apellidos']} - Cedula: {emp['cedula']}", level=1)
            reporte:Reporte = Reporte(id_empleado=emp['id_empleado'],fecha_generacion=datetime.datetime.now(),tipo_reporte=','.join(tipoReporte),contenido=','.join(tipoReporte))
            
            resultHistoria = self.insertarReporte(reporte)
            print(datetime.datetime.now())
            print(resultHistoria)
            
            if emp['departamento']:
                doc.add_paragraph(f"Departamento: {emp['departamento'].nombre}")
            if emp['rol']:
                doc.add_paragraph(f"Rol: {emp['rol'].nombre}")

            if emp['asistencias']:
                doc.add_paragraph("Asistencias:")
                for a in emp['asistencias']:
                    doc.add_paragraph(f"{a.fecha} - {a.estado_asistencia}", style='List Bullet')

            if emp['justificaciones']:
                doc.add_paragraph("Justificaciones:")
                for j in emp['justificaciones']:
                    doc.add_paragraph(f"{j.fecha} - {j.motivo} - {j.descripcion}", style='List Bullet')

            if emp['permisos']:
                doc.add_paragraph("Permisos:")
                for p in emp['permisos']:
                    doc.add_paragraph(f"{p.fecha_inicio} a {p.fecha_fin} - {p.tipo}", style='List Bullet')

            doc.add_paragraph("")  # Espacio entre empleados

        doc.save(rutaFichero)
        os.startfile(rutaFichero)
        return {'success': True, 'message': 'Se creó el reporte exitosamente.'}
  
    #crea el reporte en formato csv
    def __generar_reporte_csv_por_lotes(self, rutaFichero="reporte.csv", limit=1000, **filtros):
        
        with open(rutaFichero, mode='w', newline='', encoding='utf-8-sig') as file:
            writer = csv.writer(file,delimiter=';')
            writer.writerow([
                'Nombre', 'Apellidos', 'Cédula', 
                'Departamento', 'Rol',
                'Justificación', 'Asistencia', 'Permiso'
            ])

            for lote in self.__obtener_datos_en_lotes(limit=limit, **filtros):
                if isinstance(lote, dict) and 'success' in lote:
                    return lote
                for empleado in lote:
                    justificacion   = f"{empleado['justificacion'].fecha} - {empleado['justificacion'].motivo} - {empleado['justificacion'].descripcion}" if 'justificacion' in empleado else ''
                    asistencia      = f"{empleado['asistencia'].fecha} - {empleado['asistencia'].estado_asistencia}" if 'asistencia' in empleado else ''
                    permiso         = f"{empleado['permisos'].fecha_inicio} a {empleado['permisos'].fecha_fin} - {empleado['permisos'].tipo}" if 'permisos' in empleado else ''
                    departamento    = empleado['departamento'].nombre if 'departamento' in empleado else ''
                    rol             = empleado['rol'].nombre if 'rol' in empleado else ''
                    
                    writer.writerow([
                        empleado['nombre'],
                        empleado['apellidos'],
                        empleado['cedula'],
                        departamento,
                        rol,
                        justificacion,
                        asistencia,
                        permiso
                    ])
        
        os.startfile(rutaFichero)
        return {'success':True,'message':'Se creo el reporte exitosamente.'}